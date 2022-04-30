# 네이버 뉴스 링크 크롤링 완성
# 네이버 일반 뉴스 본문 크롤링 완성
# 네이버 연예 뉴스 본문 크롤링 완성
# 네이버 스포츠 뉴스 본문 크롤링 완성
# 검색어 입력 기능 추가
# 페이지 수 입력 기능 추가
# MS Word에 파일로 결과 저장
# MS Excel에 파일로 결과 저장
# 마지막 페이지 처리
# 2022-04-30 작성완료

import requests
from bs4 import BeautifulSoup
import time
import pyautogui
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Alignment

# 본문 내용의 불필요한 div, p삭제
def clear_div_p():

    divs = content.select("div")
    for div in divs:
        div.decompose()

    paragraphs = content.select("p")
    for p in paragraphs:
        p.decompose()

keyword = pyautogui.prompt("검색어를 입력하세요")
lastpage = int(pyautogui.prompt("몇 페이지까지 크롤링할까요?"))

# 워드 문서 생성하기
document = Document()

# 엑셀 생성하기
wb = Workbook()

# 엑셀 시트 생성하기
ws = wb.create_sheet(keyword)

# 열 너비 조절
ws.column_dimensions['A'].width = 60
ws.column_dimensions['B'].width = 60
ws.column_dimensions['C'].width = 120

# 행 번호
row = 1

# 페이지 번호
page_num = 1

for i in range(1, lastpage * 10, 10):
    response = requests.get("https://search.naver.com/search.naver?where=news&sm=tab_jum&query={}".format(keyword))
    html = response.text
    soup = BeautifulSoup(html, "html.parser")

    print("### {}페이지 크롤링 시작 ###".format(page_num))
    articles = soup.select("div.info_group")
    for article in articles:
        try:
            link_tag = article.select_one("a.info:last-child")
            url = link_tag.attrs['href']

            response = requests.get(url, headers={'User-agent':'Mozilla/5.0'})
            html = response.text
            soup_sub = BeautifulSoup(html, 'html.parser')

            # 연예 기사
            if "entertain" in response.url:
                title = soup_sub.select_one(".end_tit")
                content = soup_sub.select_one("#articeBody")
                clear_div_p()
            # 스포츠 기사
            elif "sports" in response.url:
                title = soup_sub.select_one("h4.title") 
                content = soup_sub.select_one("#newsEndContents")
                clear_div_p()
            # 일반 기사
            else:
                title = soup_sub.select_one(".media_end_head_headline")
                content = soup_sub.select_one("#dic_area")
                clear_div_p()

            print("==========링크==========\n", url)
            print("==========제목==========\n", title.text.strip())
            print("==========본문==========\n", content.text.strip())

            # 워드에 제목, 링크, 본문 저장하기
            document.add_heading(title.text.strip(), level=0)
            document.add_paragraph(url)
            document.add_paragraph(content.text.strip())

            # 엑셀에 제목, 링크, 본문 저장하기
            ws[f'A{row}'] = url
            ws[f'B{row}'] = title.text.strip()
            ws[f'C{row}'] = content.text.strip()

            # 엑셀 자동 줄바꿈
            ws[f'C{row}'].alignment = Alignment(wrap_text=True)
            row = row + 1

            time.sleep(0.5)
        except:
            pass
        
    print("### {}페이지 크롤링 종료 ###".format(page_num))

    # 마지막 페이지 여부 확인하기
    isLastPage = soup.select_one('a.btn_next').attrs['aria-disabled']
    if isLastPage == 'true':
        print("마지막 페이지 입니다.")
        break
    else:
        page_num += 1

# 워드 문서 저장하기
document.save("{}_result.docx".format(keyword))

# 엑셀 문서 저장하기
wb.save(f'{keyword}_result.xlsx')